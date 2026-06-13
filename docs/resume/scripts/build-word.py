#!/usr/bin/env python3
"""Build Word CV files from markdown sources in docs/resume/markdown/."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
MD_DIR = ROOT / "markdown"
WORD_DIR = ROOT / "word"


def strip_md(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def parse_frontmatter(content: str) -> tuple[dict[str, str], str]:
    if not content.startswith("---"):
        return {}, content
    end = content.find("\n---", 3)
    if end == -1:
        return {}, content
    block = content[3:end].strip().splitlines()
    meta: dict[str, str] = {}
    for line in block:
        if ":" in line:
            key, value = line.split(":", 1)
            meta[key.strip()] = value.strip()
    body = content[end + 4 :].lstrip("\n")
    return meta, body


def set_narrow_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(strip_md(text.lstrip("- ").strip()))
    run.font.size = Pt(10)


def add_skills_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for category, skills in rows:
        row = table.add_row().cells
        row[0].text = category
        row[1].text = skills
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
            if cell is row[0]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def build_docx(md_path: Path, out_path: Path) -> None:
    content = md_path.read_text(encoding="utf-8")
    _, body = parse_frontmatter(content)
    lines = body.splitlines()

    doc = Document()
    set_narrow_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    i = 0
    skill_rows: list[tuple[str, str]] = []
    in_skills_table = False

    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(strip_md(line[2:]))
            run.bold = True
            run.font.size = Pt(18)
            i += 1
            continue

        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(strip_md(line))
            run.bold = True
            run.font.size = Pt(11)
            i += 1
            continue

        if line.startswith("## "):
            if skill_rows:
                add_skills_table(doc, skill_rows)
                skill_rows = []
                in_skills_table = False
            add_horizontal_rule(doc)
            p = doc.add_paragraph()
            run = p.add_run(strip_md(line[3:]).upper())
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(strip_md(line[4:]))
            run.bold = True
            run.font.size = Pt(10.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        if line.startswith("|") and "---" not in line:
            parts = [p.strip() for p in line.strip("|").split("|")]
            if len(parts) == 2 and parts[0].lower() not in ("category", "categoria"):
                skill_rows.append((strip_md(parts[0]), strip_md(parts[1])))
                in_skills_table = True
            i += 1
            continue

        if line.startswith("- "):
            add_bullet(doc, line)
            i += 1
            continue

        if line.startswith("*") and line.endswith("*"):
            p = doc.add_paragraph()
            run = p.add_run(strip_md(line))
            run.italic = True
            run.font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        if line.startswith("**") and "—" in line:
            p = doc.add_paragraph()
            run = p.add_run(strip_md(line))
            run.bold = True
            run.font.size = Pt(10)
            i += 1
            continue

        p = doc.add_paragraph()
        run = p.add_run(strip_md(line))
        run.font.size = Pt(10)
        if "@" in line or "linkedin" in line or "github" in line or "+55" in line:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1

    if skill_rows:
        add_skills_table(doc, skill_rows)

    WORD_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")


def main() -> int:
    targets = [
        (MD_DIR / "LucasCruz_CV_EN.md", WORD_DIR / "LucasCruz_CV_EN.docx"),
        (MD_DIR / "LucasCruz_CV_PT.md", WORD_DIR / "LucasCruz_CV_PT.docx"),
    ]
    for md_path, out_path in targets:
        if not md_path.exists():
            print(f"Missing {md_path}", file=sys.stderr)
            return 1
        build_docx(md_path, out_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
